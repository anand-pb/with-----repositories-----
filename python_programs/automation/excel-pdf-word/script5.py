import docx

d = docx.Document('F:\\python_programs\\automation\\excel-pdf-word\\demo.docx')

print(d.paragraphs)

print(d.paragraphs[0])

print(d.paragraphs[0].text)

print(d.paragraphs[1].text)

p = d.paragraphs[1]

print(p.runs)

print(p.runs[0].text)

print(p.runs[1].text)

print(p.runs[2].text)

print(p.runs[3].text)

print(p.runs[1].bold)

print(p.runs[0].bold)

print(p.runs[0].bold == None)

print(p.runs[3].italic)

print(p.runs[3].underline == True)

p.runs[3].text = 'italic and underlined.'

d.save('F:\\python_programs\\automation\\excel-pdf-word\\demo2.docx')

print(p.style)

p.style = 'Title'

d.save('F:\\python_programs\\automation\\excel-pdf-word\\demo2.docx')