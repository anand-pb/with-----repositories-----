import docx

d = docx.Document()

d.add_paragraph('this is a new paragraph')

d.add_paragraph('this is another new paragraph')

d.save('F:\\python_programs\\automation\\excel-pdf-word\\demo4.docx')

p = d.paragraphs[0]

p.add_run('this is a new run')

p.runs

p.runs[1].bold = True

d.save('F:\\python_programs\\automation\\excel-pdf-word\\demo5.docx')
